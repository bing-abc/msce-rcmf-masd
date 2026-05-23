from __future__ import annotations

from docx import Document
from docx.shared import Pt

from _submission_utils import ROOT, write_text


COVER_DIR = ROOT / "submission_final" / "cover_letter"


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_paragraphs(document: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        document.add_paragraph(text)


def build_cover_letter() -> None:
    date_line = "May 3, 2026"
    paragraphs = [
        "Dear Editor,",
        (
            'Please consider our research article, "Multimodal Molecular Representation Learning for Polymer Glass '
            'Transition Temperature Prediction," for publication in Materials & Design.'
        ),
        (
            "This manuscript addresses polymer glass-transition temperature prediction as a structure-property problem "
            "at the computer-materials interface. We believe the work fits Materials & Design because it combines "
            "polymer informatics, multimodal molecular representation learning, and design-relevant interpretation for "
            "data-driven polymer screening while remaining anchored to a fixed and auditable evaluation protocol."
        ),
        (
            "Across 100 frozen runs, the proposed model reduced primary-test MAE from 24.67 K to 23.98 K relative to "
            "the strongest multimodal baseline. Its clearest value appeared on a fixed baseline-defined difficult subset, "
            "where MAE decreased from 29.38 K to 25.15 K. The external holdout was never used for training, validation, "
            "model selection, or trisoup selection; on that moderate chemistry-space-shift benchmark, the model showed a "
            "bounded average improvement over the strongest multimodal baseline, while a matched five-seed polyBERT audit "
            "remained slightly better on the external average."
        ),
        (
            "We have tried to present the study conservatively and transparently. The hard subgroup is defined by a fixed "
            "baseline error mask rather than by post hoc selection from the proposed model. A matched five-seed polyBERT "
            "audit remained competitive on the primary and external averages, while its difficult-subset behavior differed "
            "from the proposed multimodal framework. We therefore interpret the audit conservatively as supporting the "
            "error-tail focus of the study rather than as a definitive architecture-level comparison. The ablation results "
            "do not support RCMF as a standalone source of accuracy gain, so it is discussed only as an auxiliary "
            "diagnostic conditioning component within the full chain. Failure families, including amide-, imide-like, and "
            "heterogeneous other polymers, are disclosed explicitly rather than hidden."
        ),
        (
            "We intentionally avoid claims of universal external generalization or de novo polymer discovery; the design "
            "discussion is retrospective and bounded to the evaluated chemistry space."
        ),
        (
            "This submission is original, is not under consideration elsewhere, and has been approved by all authors. "
            "The manuscript includes declarations for competing interests, funding, acknowledgements, CRediT authorship, "
            "data availability, and generative-AI-assisted manuscript preparation. Code, processed split definitions, "
            "fixed hard-subgroup masks, result exports, statistical-test outputs, figure source data, supplementary "
            "information, highlights, and graphical abstract files are prepared for submission and release."
        ),
        (
            "A DOI-ready archival release package has been prepared; the DOI will be inserted before submission if minted "
            "in time or provided upon acceptance."
        ),
        "Thank you for your consideration.",
        "Sincerely,",
        "Peng Wang",
        "Corresponding Author",
        "School of Computer Science and Technology, Changchun University of Science and Technology",
        "Changchun 130022, China",
        "wangpeng@cust.edu.cn",
    ]

    md_lines = [date_line, "", "Editor-in-Chief", "Materials & Design", ""] + paragraphs
    write_text(COVER_DIR / "cover_letter_materials_design.md", "\n".join(md_lines) + "\n")
    write_text(COVER_DIR / "cover_letter_materials_design_final.md", "\n".join(md_lines) + "\n")
    write_text(COVER_DIR / "cover_letter_materials_design.txt", "\n".join(md_lines) + "\n")

    document = Document()
    set_default_font(document)
    document.add_paragraph(date_line)
    document.add_paragraph("Editor-in-Chief")
    document.add_paragraph("Materials & Design")
    document.add_paragraph("")
    add_paragraphs(document, paragraphs)
    document.save(str(COVER_DIR / "cover_letter_materials_design.docx"))


def build_conflict_statement() -> None:
    text = (
        "The authors declare that they have no known competing financial interests or personal relationships that could "
        "have appeared to influence the work reported in this paper."
    )
    write_text(COVER_DIR / "conflict_of_interest_statement.md", text + "\n")
    write_text(COVER_DIR / "conflict_of_interest_statement.txt", text + "\n")

    document = Document()
    set_default_font(document)
    document.add_paragraph("Conflict of Interest Statement")
    document.add_paragraph(text)
    document.save(str(COVER_DIR / "conflict_of_interest_statement.docx"))


def main() -> None:
    COVER_DIR.mkdir(parents=True, exist_ok=True)
    build_cover_letter()
    build_conflict_statement()


if __name__ == "__main__":
    main()
